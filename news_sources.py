# -*- coding: utf-8 -*-
"""
Created on Sat Jul 13 15:51:42 2019
''' sends news updates '''
@author: rayde
"""
from newspaper import Article, nlp
from GoogleNews import GoogleNews
import pandas as pd
from docx import Document
import datetime
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from string import Template
import os
import sys

class NewsSearch:
    def __init__(self, excel="C:\\Users\\rayde\\Documents\\CA_motleyfool.xlsx"):
        self.excel = excel
        self.count = 0
        self.x = 0
        
class ticker_df(NewsSearch):     
    def get_df(self):
        stocks = pd.read_excel(self.excel, usecols=[1])
        stocks.dropna(inplace=True)
        stocks = pd.DataFrame(stocks, index= None)
        return stocks

    def run_search(self, stocks):
        googlenews = GoogleNews()
        links = []
        for i, j in stocks.itertuples():
            googlenews.search(j)
            results = googlenews.getlinks()
            for link in results:
                links.append(link)
        return links
    
class articles(ticker_df):
    def build_articles(self, links):
        '''create document'''
        document = Document()
        while self.x < 10:
            for link in links:
                    try:
                        article = Article(link)
                        article.build()
                        self.content(article, document)
                        self.format_content(article, document)
                        self.write_doc(article, document)
                        self.x +=1
                    except:
                        self.count +=1
                        links.remove(link)
                        continue
        print(self.count)
        print(self.x)
        return links
        
    def content(self, article, document):
        article.nlp()
        '''create content'''
        self.title = document.add_heading(article.title,1)
        self.authors = document.add_heading(article.authors,2)
        self.date = document.add_paragraph(article.publish_date)
        self.key = document.add_paragraph(article.keywords)
        document.add_paragraph(article.summary)
        
    def format_content(self, article, document):
        '''Format Content'''
        run = self.date.add_run()
        run.font.italic =True
        run = self.authors.add_run()
        run.font.bold = True
        run = self.title.add_run()
        run.font.bold = True
        run = self.key.add_run()
        run.font.italic = True
    
    def write_doc(self, article, document):        
        '''write to file with date in name'''
        x = datetime.datetime.now()
        x = x.strftime("%Y-%m-%d")
        self.name = 'C:\\Users\\rayde\\newspapers\\{0}_news.docx'
        document.save(self.name.format(x))
        return self.name
    
class email(articles):    
    def read_template(self, filename='2019-07-13_news.docx'):
        with open(filename, 'r', encoding='utf-8') as template_file:
            template_file_content = template_file.read()
            #object_a =Template(template_file_content)
            return template_file_content
        
    def compose(self):
        x = datetime.datetime.now()
        x = x.strftime("%Y-%m-%d")
        name = '{0}_news.docx'
        name = name.format(x)
        name = "C:\\Users\\rayde\\newspapers\\" + name
        self.attachment = open(name, 'rb')
        message = self.read_template()
        msg = MIMEMultipart()
        msg['From'] = 'raydebra89@gmail.com'
        msg['To'] = 'debraray89@icloud.com'
        msg['Subject']= 'Here is your news update'
        msg.attach(MIMEText(message, 'plain'))
        msg.attach('2019-07-13_news.docx')
        self.msg = msg
        return msg
    
    def send_email(self):
        s = smtplib.SMTP('smtp.gmail.com', 587)
        s.ehlo()
        s.starttls()
        s.login('raydebra89@gmail.com', 'password')  
        s.send_message(self.compose())
        s.quit()
        
